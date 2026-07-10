from pathlib import Path
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"D:\photo\康贝儿")
SRC = ROOT / "Sanity后台管理操作指南2.0.md"
OUT = ROOT / "Sanity后台管理操作指南2.0_直观Word版.docx"

NAVY = "17324D"; TEAL = "178F8B"; PALE = "EAF6F5"; LIGHT = "F3F6F8"
GRAY = "64748B"; TEXT = "263442"; GOLD = "D89B3C"; WHITE = "FFFFFF"; RED = "A33A3A"

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn('w:shd'))
    if shd is None: shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def margins(cell, top=100, start=120, bottom=100, end=120):
    tcPr = cell._tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for k,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        e=tcMar.find(qn('w:'+k))
        if e is None: e=OxmlElement('w:'+k); tcMar.append(e)
        e.set(qn('w:w'),str(v)); e.set(qn('w:type'),'dxa')

def set_repeat(row):
    trPr=row._tr.get_or_add_trPr(); e=OxmlElement('w:tblHeader'); e.set(qn('w:val'),'true'); trPr.append(e)

def set_font(run, size=None, bold=None, color=TEXT, name='Microsoft YaHei'):
    run.font.name=name; run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'),name)
    run._element.rPr.rFonts.set(qn('w:ascii'),name); run._element.rPr.rFonts.set(qn('w:hAnsi'),name)
    if size: run.font.size=Pt(size)
    if bold is not None: run.bold=bold
    run.font.color.rgb=RGBColor.from_string(color)

def add_runs(p, text, size=9.5, color=TEXT):
    parts=re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if not part: continue
        bold=part.startswith('**'); code=part.startswith('`')
        val=part[2:-2] if bold else (part[1:-1] if code else part)
        r=p.add_run(val); set_font(r,size,bold or None,NAVY if bold else color,'Consolas' if code else 'Microsoft YaHei')

def add_toc(doc):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    r=p.add_run('目录'); set_font(r,22,True,NAVY); p.paragraph_format.space_after=Pt(12)
    info=doc.add_paragraph('以下目录打开即可查看，无需右键更新。')
    info.style='Note Box'
    sections=[
        ('01','网站设置'),('02','联系方式'),('03','布局配置'),('04','Banner 管理'),
        ('05','页面配置'),('06','服务管理'),('07','优势管理'),('08','案例管理'),
        ('09','科普管理'),('10','常见问题'),('11','咨询线索'),
        ('附录 A','图片尺寸速查表'),('附录 B','快速修改指南'),('附录 C','注意事项'),
    ]
    table=doc.add_table(rows=0,cols=2); table.autofit=False
    for num,title in sections:
        cells=table.add_row().cells; cells[0].text=num; cells[1].text=title
    for i,row in enumerate(table.rows):
        for j,c in enumerate(row.cells):
            c.width=Inches(.95 if j==0 else 5.85); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; margins(c,85,120,85,120)
            if i%2: shade(c,LIGHT)
            for p2 in c.paragraphs:
                p2.paragraph_format.space_after=Pt(0)
                for rr in p2.runs: set_font(rr,9.5,j==0,TEAL if j==0 else TEXT)

def setup(doc):
    sec=doc.sections[0]; sec.page_width=Inches(8.27); sec.page_height=Inches(11.69)
    sec.top_margin=Inches(.72); sec.bottom_margin=Inches(.7); sec.left_margin=Inches(.72); sec.right_margin=Inches(.72)
    styles=doc.styles
    n=styles['Normal']; n.font.name='Microsoft YaHei'; n._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); n.font.size=Pt(9.5); n.font.color.rgb=RGBColor.from_string(TEXT)
    n.paragraph_format.space_after=Pt(5); n.paragraph_format.line_spacing=1.18
    for name,size,color,before,after in [('Title',28,NAVY,0,8),('Subtitle',13,GRAY,0,8),('Heading 1',18,NAVY,16,8),('Heading 2',14,TEAL,12,6),('Heading 3',11,NAVY,9,4)]:
        s=styles[name]; s.font.name='Microsoft YaHei'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); s.font.size=Pt(size); s.font.bold=name!='Subtitle'; s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
    for name,fill,color in [('Note Box',PALE,TEAL),('Warning Box','FFF4E5',RED)]:
        if name not in styles: styles.add_style(name,WD_STYLE_TYPE.PARAGRAPH)
        s=styles[name]; s.font.name='Microsoft YaHei'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); s.font.size=Pt(9.5); s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.left_indent=Inches(.16); s.paragraph_format.right_indent=Inches(.12); s.paragraph_format.space_before=Pt(5); s.paragraph_format.space_after=Pt(7)
        pPr=s._element.get_or_add_pPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); pPr.append(shd)
    header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT; rr=header.add_run('SANITY 后台管理 · 操作指南 2.0'); set_font(rr,8,True,GRAY)
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rr=footer.add_run('康贝儿网站内容管理  ·  '); set_font(rr,8,False,GRAY)
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); footer._p.append(fld)

def cover(doc):
    for _ in range(5): doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('SANITY'); set_font(r,12,True,TEAL); p.paragraph_format.space_after=Pt(16)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('后台管理操作指南'); set_font(r,30,True,NAVY); p.paragraph_format.space_after=Pt(8)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('2.0'); set_font(r,18,True,GOLD)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(18); r=p.add_run('网站内容维护 · 页面配置 · 图片规范 · 咨询线索'); set_font(r,11,False,GRAY)
    for _ in range(6): doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('最后更新：2025-01'); set_font(r,9,False,GRAY)
    doc.add_page_break()

def add_quick_start(doc):
    doc.add_heading('使用前先看',1)
    p=doc.add_paragraph('这份手册适合日常内容维护人员。先从下方路径找到对应模块，再按字段说明填写；涉及图片时，请优先查看文末“图片尺寸速查表”。'); p.style='Note Box'
    items=[('改电话','网站设置 / 联系方式 / 咨询表单配置'),('改微信','联系方式 / 网站设置 / 布局配置'),('改页面标题','Banner 管理 / 页面配置'),('改底部咨询区','网站设置 → CTA 相关字段'),('查看用户留言','咨询线索')]
    t=doc.add_table(rows=1,cols=2); t.style='Table Grid'; t.autofit=False
    t.columns[0].width=Inches(1.6); t.columns[1].width=Inches(5.2)
    t.cell(0,0).text='我要做什么'; t.cell(0,1).text='去哪里修改'
    for a,b in items:
        c=t.add_row().cells; c[0].text=a; c[1].text=b
    style_table(t,[2200,7160])
    doc.add_page_break(); add_toc(doc); doc.add_page_break()

def style_table(t,widths=None):
    t.autofit=False
    for i,row in enumerate(t.rows):
        for j,c in enumerate(row.cells):
            c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; margins(c)
            if widths: c.width=Inches(widths[j]/1440)
            if i==0: shade(c,NAVY)
            elif i%2==0: shade(c,LIGHT)
            for p in c.paragraphs:
                p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.05
                for r in p.runs: set_font(r,8.3,i==0,WHITE if i==0 else TEXT)
        if i==0: set_repeat(row)

def parse(doc, text):
    lines=text.splitlines(); i=0; skipped_toc=False
    while i<len(lines):
        line=lines[i].rstrip()
        if line.startswith('# '): i+=1; continue
        if line=='## 📑 目录':
            skipped_toc=True; i+=1
            while i<len(lines) and not lines[i].startswith('## 1.'): i+=1
            continue
        if not line or line=='---': i+=1; continue
        m=re.match(r'^(#{2,4})\s+(.*)',line)
        if m:
            level=min(len(m.group(1))-1,3); title=re.sub(r'^[📸🎯⚠️]\s*','',m.group(2))
            doc.add_heading(title,level); i+=1; continue
        if line.startswith('> '):
            p=doc.add_paragraph(); p.style='Note Box'; add_runs(p,'提示：'+line[2:]); i+=1; continue
        if line.startswith('|') and i+1<len(lines) and re.match(r'^\|[\s:|-]+\|$',lines[i+1]):
            rows=[]
            while i<len(lines) and lines[i].startswith('|'):
                if not re.match(r'^\|[\s:|-]+\|$',lines[i]): rows.append([x.strip() for x in lines[i].strip('|').split('|')])
                i+=1
            cols=max(len(r) for r in rows); t=doc.add_table(rows=len(rows),cols=cols); t.style='Table Grid'
            for ri,row in enumerate(rows):
                for ci,val in enumerate(row):
                    t.cell(ri,ci).text=''; add_runs(t.cell(ri,ci).paragraphs[0],val,8.3,WHITE if ri==0 else TEXT)
            if cols==4: widths=[1900,2500,3160,1800]
            elif cols==3: widths=[2300,4100,2960]
            elif cols==2: widths=[2600,6760]
            else: widths=[9360//cols]*cols
            style_table(t,widths); doc.add_paragraph().paragraph_format.space_after=Pt(1); continue
        if re.match(r'^\d+\.\s+',line):
            p=doc.add_paragraph(style='List Number'); add_runs(p,re.sub(r'^\d+\.\s+','',line)); i+=1; continue
        if line.startswith('- '):
            p=doc.add_paragraph(style='List Bullet'); add_runs(p,line[2:]); i+=1; continue
        if line.startswith('**后台位置：**'):
            p=doc.add_paragraph(); p.style='Note Box'; add_runs(p,line); i+=1; continue
        p=doc.add_paragraph(); add_runs(p,line); i+=1

def main():
    doc=Document(); setup(doc); cover(doc); add_quick_start(doc)
    text=SRC.read_text(encoding='utf-8'); parse(doc,text)
    doc.core_properties.title='Sanity 后台管理操作指南 2.0'; doc.core_properties.subject='康贝儿网站后台内容维护手册'; doc.core_properties.author='康贝儿'
    doc.save(OUT); print(OUT)

if __name__=='__main__': main()
